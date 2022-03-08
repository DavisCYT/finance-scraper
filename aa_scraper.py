#!/usr/bin/env python
# coding: utf-8

# In[116]:


from selenium import webdriver
from bs4 import BeautifulSoup
import pandas as pd
from selenium.webdriver.edge.service import Service as EdgeService
from selenium.webdriver.common.by import By
from webdriver_manager.microsoft import EdgeChromiumDriverManager
import requests
import json
import math
import re
import numpy as np
import matplotlib.pyplot as plt
import csv
import docx
import requests
driver = webdriver.Edge(EdgeChromiumDriverManager().install())


# ## Testing Parameters

# In[117]:


stock = '8055'
lang = 'en'
report = 'cf'
url = f'http://www.aastocks.com/{lang}/stocks/analysis/company-fundamental/profit-loss?symbol={stock}'


# # AAStock Scraper

# ## AAStock Financial Scraper

# In[118]:


def aa_fins(stock, lang = 'en', report = 'annual', fins = 'pl', num_yr = 5):
    if isinstance(stock, int):
        stock = str(stock)
    if report == 'annual':
        period = '4'
    if report == 'interim':
        period = '2'
    if report == 'quarterly':
        period = '0'
    
    if fins == 'pl':
        fins = 'profit-loss'
        xpath = '//tr[contains(@ref, "PL")]'
    if fins == 'bs':
        fins = 'balance-sheet'
        xpath = '//tr[contains(@ref, "BS")]'
    if fins == 'cf':
        fins = 'cash-flow'
        xpath = '//tr[contains(@ref, "CF")]'
    if fins == 'fr':
        fins = 'financial-ratios'
        xpath = '//tr[contains(@ref, "FR")]'
        
    url = f'http://www.aastocks.com/{lang}/stocks/analysis/company-fundamental/{fins}?symbol={stock}&period={period}'
    driver.get(url)
    content = driver.page_source
    soup = BeautifulSoup(content)

    dic = {}
    
    for tr in driver.find_elements(By.XPATH, xpath):
        profile = []
        for td in tr.find_elements(By.TAG_NAME, 'td'):
            profile.append(td.text)
            dic[profile[0]] = profile[1:6]

    df = pd.DataFrame.from_dict(dic, orient = 'index')
    df.columns = df.iloc[0]
    df = df.drop(list(dic)[0])
    df = df[df.columns[::-1]]
    # Add one more level of Column
    if lang == 'en':
        report = report.capitalize()
    if lang == 'tc' or lang == 'sc':
        if report == 'annual':
            report = '年報'
        elif report == 'interim':
            report = '期中'
        else:
            report = '季報'
        
    fs = [report.capitalize() for i in range(0, len(df.columns))]
    columns = tuple(zip(fs, df.columns))
    df.columns = pd.MultiIndex.from_tuples(columns)
    df = df.iloc[:, :num_yr]
    return df


# In[329]:


df = aa_fins(8045, report = 'interim', num_yr = 3, fins = 'pl')


# ## AAStock Profile Scraper

# In[234]:


def aa_profile(stock, lang = 'en'):
        
    url = f'http://www.aastocks.com/{lang}/stocks/analysis/company-fundamental/company-profile?symbol={stock}'
    driver.get(url)
    content = driver.page_source
    soup = BeautifulSoup(content)

    xpath = '//table[contains(@class, ("cnhk-cf tblM s4 s5 mar15T"))]/tbody/tr'
    title_xpath = '//div[contains(@class, ("ns1 white mar10T"))]/div[contains(@class, ("title"))]'
    col = driver.find_element(By.XPATH, title_xpath).text
    dic = {}

    for tr in driver.find_elements(By.XPATH, xpath):
        profile = []
        for td in tr.find_elements(By.TAG_NAME, 'td'):
            profile.append(td.text)
        dic[profile[0]] = profile[1]
    
    df = pd.DataFrame.from_dict(dic, orient = 'index')
    df.rename(columns = {0 : col}, inplace = True)
    
    return df


# In[316]:


def bus_sum(stock, lang = 'en'):
    cor_pro = aa_profile(stock, lang).iloc[5][0]
    bus = cor_pro.split('\n')[1]
    bus_title = cor_pro.split('\n')[0].strip(':')
    return bus_title, bus


# In[317]:


b_title, bus = bus_sum(8045)


# In[318]:


aa_profile(8045)


# ## AAStock Corporate Information Scraper

# In[229]:


def aa_cor_info(stock, lang = 'en'):
        
    url = f'http://www.aastocks.com/{lang}/stocks/analysis/company-fundamental/company-information?symbol={stock}'
    driver.get(url)
    content = driver.page_source
    soup = BeautifulSoup(content)

    xpath = '//table[contains(@class, ("cnhk-cf tblM s4 s5 mar15T"))]/tbody/tr'
    title_xpath = '//div[contains(@class, ("ns1 white mar10T"))]/div[contains(@class, ("title"))]'
    col = driver.find_element(By.XPATH, title_xpath).text
    dic = {}

    for tr in driver.find_elements(By.XPATH, xpath):
        profile = []
        for td in tr.find_elements(By.TAG_NAME, 'td'):
            profile.append(td.text)
        dic[profile[0]] = profile[1]
    
    df = pd.DataFrame.from_dict(dic, orient = 'index')
    df.rename(columns = {0 : col}, inplace = True)
    
    return df


# In[191]:


def shareholders(stock, lang = 'en'):
    shareholder = aa_cor_info(stock, lang).loc['Substantial Shareholders'][0].split('\n')
    shareholder_dic = {'Substantial Shareholders' : shareholder}
    return shareholder_dic


# In[230]:


df = aa_cor_info(8045)


# ## AAStock Basic Information Scraper

# In[232]:


def aa_basic(stock, lang = 'en'):
    url = f'http://www.aastocks.com/{lang}/stocks/analysis/company-fundamental/basic-information?symbol={stock}'
    driver.get(url)
    content = driver.page_source
    soup = BeautifulSoup(content)

    xpath = '//table[contains(@class, ("cnhk-cf tblM s4 s5 mar15T"))]/tbody/tr'
    title_xpath = '//div[contains(@class, ("ns1 white mar10T"))]/div[contains(@class, ("title"))]'
    col = driver.find_element(By.XPATH, title_xpath).text
    dic = {}

    for tr in driver.find_elements(By.XPATH, xpath):
        profile = []
        for td in tr.find_elements(By.TAG_NAME, 'td'):
            profile.append(td.text)
        dic[profile[0]] = profile[1]
    
    df = pd.DataFrame.from_dict(dic, orient = 'index')
    df.rename(columns = {0 : col}, inplace = True)
    return df


# In[233]:


aa_basic(8045)


# # Yorkshire Profile AAStock Scraper

# ### Title (Header) : Company Name + Stock Code (
# ### Profile (Paragraph) : Business Summary (aa_cor_pro)
# ### Snapshot (Table) : Share Price, Market Cap, Share Out., PE, Listing Date (aa_basic)
# ### Corporate Profile (Table) : Substantial Shareholders, Directors (aa_cor_info)
# ### Financial (Table) : Balance Sheet, Profit & Loss (aa_fins)

# In[305]:


df = aa_basic(8045,'tc')
basic_list = ['Name', 'Close Price', 'Market Capital', 'Issued Capital(share)', 'PE', 'Listing Date']


# In[304]:


pd.DataFrame(list(df.index))
[0, 30, 25, 9, 14, 3]


# In[308]:


aa_cor_info(8045)


# In[313]:


def ys_aa_profile(stock, lang = 'en'):

#     pl_list = ['Total Turnover', 'Gross Profit', 'Operating Profit', 'Net Profit', 'EBITDA']
#     bs_list = ['Cash On Hand', 'Total Assets', 'Total Debt', 'Intangible Assets', 'Minority interests', 'Net Assets']
#     basic_list = ['Name', 'Close Price', 'Market Capital', 'Issued Capital(share)', 'PE', 'Listing Date']
#     cor_list = ['Substantial Shareholders', 'Directors']
    
    # Profile and Loss Statement
    pl_index = [0, 4, 6, 12,15]
    an_3yr_pl = aa_fins(stock, lang, report = 'annual', fins = 'pl', num_yr = 3).iloc[pl_index]
    int_pl = aa_fins(stock, lang, report = 'interim', fins = 'pl', num_yr = 1).iloc[pl_index]
    pl = pd.concat([int_pl, an_3yr_pl], axis = 1)
#     pl.rename(index = {'Total Turnover' : 'Revenue' }, inplace = True)
    # Balance Sheet
    bs_index = [5, 9, 27, 29, 23, 28]
    bs = aa_fins(stock, lang, report = 'interim', fins = 'bs').iloc[bs_index].iloc[:, :1]  
    # Business Summary Title and Paragraph
    bus_title, bus  = bus_sum(stock, lang)
    # Basic Info of Stock
    basic_index = [0, 30, 25, 9, 14, 3]
    basic = aa_basic(stock, lang).iloc[basic_index]
#     basic.loc['Issued Capital(share)'] = aa_profile(stock).loc['Share Issued (share)']
#     basic = basic.rename(index = {'Close Price' : 'Price',
#                                         'Market Capital' : 'Market Cap.',
#                                         'Issued Capital(share)' : 'Share Out.'})
    # Company Name
    name = aa_basic(stock, lang).iloc[0][0]
    # Corporate Structure
    info = aa_cor_info(stock, lang).iloc[0:2]
# ys_title : 'Business Summary'
# ys_bus : Paragraph of Business Summary
# ys_basic : Basic Info,
# ys_bs : Balance Sheet
# ys_info : Corporate Structure
# ys_pl : Profile and Loss
    return name , bus_title, bus , basic, bs, pl, info


# In[237]:


name , bus_title, bus , basic, bs, pl, info = ys_aa_profile(8045)


# In[341]:


x = aa_basic(8045, 'tc')


# In[342]:


x.iloc[0][0]


# In[188]:


info


# # Dataframe to Excel File

# ## Skills : Export Dataframe to Excel File

# In[155]:


def to_xlsx(stock, lang = 'en'):
    ys_title, ys_bus, ys_basic, ys_bs, ys_info, ys_pl = ys_aa_profile(stock, lang)
    file_name = 'Stock_'+ str(stock) +'.xlsx'
    xlwriter = pd.ExcelWriter(file_name)
    
    ys_basic.to_excel(xlwriter, sheet_name = 'Basic', index = True)
    ys_info.to_excel(xlwriter, sheet_name = 'Corporate Info', index = True)
    ys_bs.to_excel(xlwriter, sheet_name = 'Balance Sheet', index = True)
    ys_pl.to_excel(xlwriter, sheet_name = 'Profit_Loss', index = True)
    
    xlwriter.close()


# # Company Profile Generator

# ## Add Heading and Paraphragh

# In[343]:


from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.run import Font, Run, WD_BREAK 
from docx.shared import Pt


# In[349]:


def add_title_bus(name, bus, bus_title,doc, doc_name):

# Add the document title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(name)
    run.font.size = Pt(18)
    run.bold = True

# Add the Business Summary Heading
    bs_head = doc.add_paragraph()
    bs_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bs_head.add_run(bus_title).bold = True    

# Add the Business Summary Paragraph
    bs_para = doc.add_paragraph(bus)
    
    


# In[347]:


doc = docx.Document()
add_title_bus('Hello','Haha', '集團業務',doc, 'DocName')
doc.save('DocName')


# ## Add Table

# In[350]:


def add_table(df, doc, doc_name):
    
    t = doc.add_table(df.shape[0]+1, df.shape[1] + 1, style = 'Light Grid Accent 3')
    # Create the Column Row
    for col in range(df.shape[-1]):
        t.cell(0, col +1).text = df.columns[col]

    # Create the Index Column
    for index in range(df.shape[0]):
        t.cell(index+1, 0).text = df.index[index]

    # add the rest of the data frame
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            t.cell(i+1,j+1).text = str(df.values[i,j])
    
    doc.add_paragraph().add_run().add_break()


# In[326]:


doc = docx.Document()


# # Stock Profile Creator

# In[351]:


def profile_doc(doc_name, stock, lang = 'en'):
    
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.text.run import Font, Run, WD_BREAK
    from docx.shared import Pt
    
    doc_name += '.docx'
    doc = docx.Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)
    
    name , bus_title, bus , basic, bs, pl, info = ys_aa_profile(stock, lang)
    
    add_title_bus(name, bus, bus_title, doc, doc_name)
    add_table(basic, doc, doc_name)
    add_table(info, doc, doc_name)
    doc.add_paragraph().add_run().add_break(break_type = WD_BREAK.PAGE)
    add_table(bs, doc, doc_name)
    add_table(pl, doc, doc_name)
    
    doc.save(doc_name)


# In[352]:


profile_doc('Test2', 8045)


# In[353]:


profile_doc('8045', 8045, 'tc')


# In[ ]:




